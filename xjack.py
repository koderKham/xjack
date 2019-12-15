from docx import Document
from docxtpl import DocxTemplate
import models

def generate():
    doc = DocxTemplate("NG_template.docx")
    dig = models.Materials(def_name, case_number, judge, date_of_today)
    doc.render(dig.context)
    doc.save("{}'s NG Plea.docx".format(dig.def_name))
    doc = DocxTemplate("d4d.docx")
    dig = models.Materials(def_name, case_number, judge, date_of_today)
    doc.render(dig.context)
    doc.save("{}'s Disco Demand.docx".format(dig.def_name))
    doc = DocxTemplate("NOA.docx")
    dig = models.Materials(def_name, case_number, judge, date_of_today)
    doc.render(dig.context)
    doc.save("{}'s NOA.docx".format(dig.def_name))

menu_selection = raw_input("Would you like to generate your usual criminal case documents? ")
if menu_selection == "y":
    def generate():
        def_name = raw_input("What is the defendant's name? ")
        case_number = raw_input("What is the case number? ")
        judge = raw_input("Whats the judge's name? ")
        date_of_today = raw_input("What is today's date? ")
        doc = DocxTemplate("NG_template.docx")
        dig = models.Materials(def_name, case_number, judge, date_of_today)
        doc.render(dig.context)
        doc.save("{}'s NG Plea.docx".format(dig.def_name))
        doc = DocxTemplate("d4d.docx")
        dig = models.Materials(def_name, case_number, judge, date_of_today)
        doc.render(dig.context)
        doc.save("{}'s Disco Demand.docx".format(dig.def_name))
        doc = DocxTemplate("NOA.docx")
        dig = models.Materials(def_name, case_number, judge, date_of_today)
        doc.render(dig.context)
        doc.save("{}'s NOA.docx".format(dig.def_name))
    generate()

def make_client():
    clientName = raw_input("Whats the client's name? ")
    clientAddress = raw_input("Whats the client's address? ")
    clientCtyStZip = raw_input("City, State, and Zip? ")
    clientEmail = raw_input("Email? ")
    date = raw_input("Whats the date? ")
    models.Client(clientName, date, clientAddress, clientCtyStZip, clientEmail)








'''context = { 'date' : date, 'clientName': clientName, "clientCtyStZip" : clientCtyStZip,
"clientFee" : clientFee, "lawyerAssigned" : lawyerAssigned, "mattersDescription" : mattersDescription,
"clientServices" : clientServices, "estimatedTimeFrame" : estimatedTimeFrame }'''



















'''Build a fully automated document rendering program
 ****IDEAS****
 Client Engagement form

class ClientEngagementLetter:
    def __init__(self, name, date, servicesRequested):
        self.clientName = name
        self.date = date
        self.servicesRequested = servicesRequested

name = raw_input("Whats the clients name? ")
date = raw_input("Whats the date? ")
service = raw_input("Whats the service is the client requesting? ")

x = ClientEngagementLetter(name, date, service)

paragraph = document.add_paragraph(x.servicesRequested)
paragraph = document.add_paragraph(x.date)
paragraph = document.add_paragraph(x.clientName)

document.save('test.docx')
x = ClientEngagementLetter(name, date, service)
print(x.clientName)
print(x.date)
print(x.servicesRequested)'''
